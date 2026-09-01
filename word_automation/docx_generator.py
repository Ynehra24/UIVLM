from __future__ import annotations

import ast
import contextlib
import io
import os
import re
import traceback
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

import docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Inches, Mm, Pt, RGBColor


class PythonDocxCodeGenerator:
    """
    Generates and executes concise, robust python-docx code for file manipulation tasks.
    Supports creating new documents as well as referencing and editing existing documents without overwriting.
    """

    CONTEXT = r"""
# Python-DOCX Reference & Best Practices

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 1. Opening/Creating Documents:
# A) Brand NEW document:
# filename = 'new_doc.docx'
# doc = Document()

# B) EDITING an EXISTING document (Preserve existing content):
# filename = 'existing_doc.docx'
# doc = Document(filename)

# 2. Substantive Paragraphs with Specific Font, Color & Size:
# p = doc.add_paragraph()
# run = p.add_run('India is a vibrant civilization known for its profound cultural diversity, ancient history, and linguistic pluralism. From the snow-capped Himalayas in the north to the tropical coastlines of the south, it encompasses hundreds of languages, traditions, and architectural marvels. In the modern era, the nation has emerged as a global technological powerhouse and the world’s most populous democracy, driving forward advancements in space exploration, digital infrastructure, and sustainable development while steadfastly preserving its timeless heritage.')
# run.font.name = 'Impact'
# run.font.size = Pt(4)
# run.font.color.rgb = RGBColor(0, 128, 0) # Green

# 3. Saving:
# doc.save(filename)
"""

    def _extract_clean_code(self, raw_content: str) -> str:
        """Extract clean Python code from model response, stripping thinking tags and code fences."""
        if not raw_content:
            return ""

        text = raw_content.strip()
        text = re.sub(r"<think>[\s\S]*?</think>", "", text, flags=re.IGNORECASE).strip()

        if "```" in text:
            match = re.search(r"```(?:python)?\s*([\s\S]*?)(?:```|$)", text, re.IGNORECASE)
            if match:
                text = match.group(1).strip()

        lines = text.splitlines()
        code_start_keywords = (
            "import ", "from ", "#", "doc =", "doc=", "def ", "class ",
            "table =", "table=", "p =", "p=", "run =", "run=", "filename =", "filename="
        )
        start_idx = 0
        found_code = False
        for idx, line in enumerate(lines):
            l_str = line.strip()
            if any(l_str.startswith(kw) for kw in code_start_keywords):
                start_idx = idx
                found_code = True
                break

        if found_code and start_idx > 0:
            text = "\n".join(lines[start_idx:]).strip()

        return text

    def generate(
        self,
        task_description: str,
        nemotron_client: Any,
        model: str = "nvidia/nemotron-3-super-120b-a12b",
        error_context: Optional[str] = None,
    ) -> str:
        code, _ = self.generate_with_usage(task_description, nemotron_client, model, error_context)
        return code

    def generate_with_usage(
        self,
        task_description: str,
        nemotron_client: Any,
        model: str = "nvidia/nemotron-3-super-120b-a12b",
        error_context: Optional[str] = None,
    ) -> Tuple[str, int]:
        """
        Generate concise, executable python-docx code and track token usage.
        """
        system_prompt = f"""You are a python-docx expert.

Generate clean, accurate, production-ready Python code using the python-docx library.

{self.CONTEXT}

CRITICAL RULES:
1. Generate ONLY valid, executable Python code with all necessary imports.
2. PARAGRAPH DEPTH: When asked to write or add a paragraph about a topic, write a full, well-developed, substantive paragraph (4-6 comprehensive sentences with depth), not just a single brief sentence!
3. PRESERVE EXISTING FILES: When editing an existing document, load it with `doc = Document(filename)` and append/modify without wiping existing content.
4. NEW FILES: When creating a new document, use `doc = Document()`.
5. FORMATTING & COLORS: Accurately apply requested fonts (e.g. Impact), colors (RGBColor), and sizes (Pt).
6. BULLET POINTS: Put bullet text directly into `doc.add_paragraph('Bullet text', style='List Bullet')`.
7. TABLES: Use `table.cell(r, c).text = ...` and `table.style = 'Table Grid'`.
8. Always save the document at the end with `doc.save(filename)`.
9. No conversational explanations, output ONLY executable Python code."""

        user_content = task_description
        if error_context:
            user_content += f"\n\nIMPORTANT: Fix this previous error:\n{error_context}"

        tokens_used = 0
        kwargs: Dict[str, Any] = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            "temperature": 0.1,
            "max_tokens": 2048,
        }

        try:
            response = nemotron_client.chat.completions.create(
                **kwargs,
                extra_body={"chat_template_kwargs": {"enable_thinking": False}},
            )
        except Exception:
            response = nemotron_client.chat.completions.create(**kwargs)

        if hasattr(response, "usage") and response.usage:
            tokens_used = getattr(response.usage, "total_tokens", 0)

        content = response.choices[0].message.content or ""
        code = self._extract_clean_code(content)
        return code, tokens_used

    def execute(self, code: str, working_dir: Optional[str] = None) -> Tuple[bool, str, str]:
        """
        Execute generated python-docx code in a structured namespace.
        """
        if not code.strip():
            return False, "", "No code provided to execute"

        try:
            ast.parse(code)
        except SyntaxError as syn_err:
            return False, "", f"SyntaxError in generated code: {syn_err}"

        # Extract filename if present in code to provide resilient default doc
        filename_match = re.search(r"filename\s*=\s*['\"]([^'\"]+)['\"]", code)
        doc_inst = None
        if filename_match:
            fname = filename_match.group(1)
            if os.path.exists(fname):
                try:
                    doc_inst = Document(fname)
                except Exception:
                    doc_inst = Document()
            else:
                doc_inst = Document()

        # Setup namespace
        namespace: Dict[str, Any] = {
            "os": os,
            "re": re,
            "docx": docx,
            "Document": Document,
            "doc": doc_inst or Document(),
            "Inches": Inches,
            "Pt": Pt,
            "Cm": Cm,
            "Mm": Mm,
            "RGBColor": RGBColor,
            "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
            "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
            "parse_xml": parse_xml,
            "nsdecls": nsdecls,
            "Path": Path,
        }

        try:
            import openpyxl
            namespace["openpyxl"] = openpyxl
        except ImportError:
            pass

        try:
            import pptx
            namespace["pptx"] = pptx
        except ImportError:
            pass

        stdout_buf = io.StringIO()
        stderr_buf = io.StringIO()

        try:
            with contextlib.redirect_stdout(stdout_buf), contextlib.redirect_stderr(stderr_buf):
                exec(code, namespace)

            captured_out = stdout_buf.getvalue().strip()
            captured_err = stderr_buf.getvalue().strip()
            output_msg = captured_out or "Code executed successfully"
            if captured_err:
                output_msg += f"\nStderr:\n{captured_err}"
            return True, output_msg, ""
        except Exception as e:
            err_details = traceback.format_exc()
            return False, stdout_buf.getvalue().strip(), f"Execution error: {e}\n{err_details}"
