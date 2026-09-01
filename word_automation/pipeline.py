from __future__ import annotations

import os
import re
from typing import Any, Dict, List, Optional

from docx import Document
from openai import OpenAI

from .applescript_generator import AppleScriptCodeGenerator
from .config import AutomationConfig
from .decomposer import TaskDecomposer
from .docx_generator import PythonDocxCodeGenerator


class WordAutomationPipeline:
    """
    Main orchestration layer for Word automation using a 3-layer architecture:
    1. Nemotron (Task decomposition & routing)
    2. Python-docx (File-level manipulation - 70-80% of tasks)
    3. Osaurus AppleScript-8B (App-level control - edge cases, with Nemotron fallback if Osaurus offline)

    Supports:
    - Referencing and editing existing documents without overwriting.
    - Creating new documents when the referenced file does not exist.
    - Explicit overwrite when requested via parameter or prompt keyword.
    """

    OVERWRITE_PATTERNS = [
        r"\boverwrite\b",
        r"\brecreate\b",
        r"\bstart from scratch\b",
        r"\bcreate from scratch\b",
        r"\bfrom scratch\b",
        r"\breplace existing\b",
        r"\bwipe and recreate\b",
        r"\bfresh document\b",
    ]

    def __init__(
        self,
        nemotron_api_key: Optional[str] = None,
        osaurus_base_url: str = "http://localhost:8080/v1",
        config: Optional[AutomationConfig] = None,
        nemotron_base_url: Optional[str] = None,
        nemotron_client: Optional[Any] = None,
        osaurus_client: Optional[Any] = None,
        max_retries: int = 2,
    ):
        self.config = config or AutomationConfig.from_env()
        self.max_retries = max_retries

        api_key = nemotron_api_key or self.config.nemotron_api_key
        base_url = nemotron_base_url or self.config.nemotron_base_url
        os_url = osaurus_base_url or self.config.osaurus_base_url

        if nemotron_client is not None:
            self.nemotron = nemotron_client
        else:
            self.nemotron = OpenAI(
                base_url=base_url,
                api_key=api_key or "not-set",
            )

        if osaurus_client is not None:
            self.osaurus = osaurus_client
        else:
            self.osaurus = OpenAI(
                base_url=os_url,
                api_key="not-needed",
            )

        self.decomposer = TaskDecomposer()
        self.docx_gen = PythonDocxCodeGenerator()
        self.applescript_gen = AppleScriptCodeGenerator(
            osaurus_client=self.osaurus,
            fallback_client=self.nemotron,
            fallback_model=self.config.nemotron_model,
            sdef_cache_dir=self.config.cache_dir,
            timeout=self.config.applescript_timeout,
            sdef_timeout=self.config.sdef_extraction_timeout,
        )

        self.execution_log: List[Dict[str, Any]] = []

    def _determine_target_filename(self, user_request: str, explicit_file: Optional[str] = None) -> str:
        """Derive or extract a clean target filename for the document."""
        if explicit_file and explicit_file.strip():
            fname = explicit_file.strip()
            return fname if fname.lower().endswith(".docx") else f"{fname}.docx"

        match = re.search(r"[\w\-./\\]+\.docx", user_request, re.IGNORECASE)
        if match:
            return match.group(0)

        # Extract title keywords or default
        clean = re.sub(r"[^\w\s]", "", user_request.lower())
        words = [
            w for w in clean.split()
            if w not in {"write", "me", "a", "an", "the", "on", "and", "in", "page", "pages", "document", "docx", "create", "edit", "update", "modify"}
        ]
        if words:
            return "_".join(words[:4]).capitalize() + ".docx"
        return "Document.docx"

    def _is_explicit_overwrite(self, user_request: str, overwrite_flag: bool) -> bool:
        """Check if overwrite was explicitly requested via flag or prompt keywords."""
        if overwrite_flag:
            return True
        req_lower = user_request.lower()
        for pattern in self.OVERWRITE_PATTERNS:
            if re.search(pattern, req_lower):
                return True
        return False

    def _get_document_summary(self, filename: str) -> str:
        """Extract a structured summary of an existing document for model context."""
        try:
            doc = Document(filename)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables_count = len(doc.tables)
            sections_count = len(doc.sections)
            summary = (
                f"Existing document '{filename}': {len(paragraphs)} paragraph(s), "
                f"{tables_count} table(s), {sections_count} section(s)."
            )
            if paragraphs:
                preview = "\n".join([f"  P{i+1}: {p[:120]}" for i, p in enumerate(paragraphs[:6])])
                summary += f"\nExisting Paragraphs Preview:\n{preview}"
            return summary
        except Exception as e:
            return f"Existing file '{filename}' could not be inspected: {e}"

    def execute(
        self,
        user_request: str,
        target_file: Optional[str] = None,
        overwrite: bool = False,
        verbose: bool = True,
    ) -> Dict[str, Any]:
        """
        Execute a Word automation request end-to-end.

        Args:
            user_request: The task prompt or instruction
            target_file: Optional explicit target/reference .docx file
            overwrite: If True, explicitly recreate/overwrite existing document.
                       If False (default), existing documents are edited and preserved.
            verbose: Whether to print progress logs to stdout
        """
        target_filename = self._determine_target_filename(user_request, target_file)
        explicit_overwrite = self._is_explicit_overwrite(user_request, overwrite)
        doc_already_exists = os.path.exists(target_filename) and os.path.isfile(target_filename)

        if doc_already_exists and not explicit_overwrite:
            mode = "EDIT_EXISTING"
        elif doc_already_exists and explicit_overwrite:
            mode = "OVERWRITE"
            try:
                os.remove(target_filename)
            except Exception:
                pass
        else:
            mode = "CREATE_NEW"

        results: Dict[str, Any] = {
            "status": "success",
            "mode": mode,
            "target_file": target_filename,
            "original_request": user_request,
            "subtasks": [],
            "errors": [],
            "total_tokens_used": 0,
            "total_tokens": 0,
        }

        if verbose:
            print(f"📋 Task: {user_request}")
            print(f"📁 Target Document: '{target_filename}' (Mode: {mode})\n")

        doc_context = None
        if mode == "EDIT_EXISTING":
            doc_context = self._get_document_summary(target_filename)
            if verbose:
                print(f"📖 Context: {doc_context}\n")

        total_tokens = 0

        try:
            if verbose:
                print("🔍 Decomposing task with Nemotron...")

            subtasks, decomp_tokens = self.decomposer.decompose_with_usage(
                user_request,
                self.nemotron,
                model=self.config.nemotron_model,
                doc_context=doc_context,
            )
            total_tokens += decomp_tokens

            if verbose:
                print(f"   Found {len(subtasks)} sub-task(s)\n")

            all_succeeded = True
            any_succeeded = False

            for i, subtask in enumerate(subtasks):
                task_type = subtask.get("type", "FILE_MANIPULATION")
                description = subtask.get("description", "")
                is_first_subtask = (i == 0)

                # Ensure consistent filename in description
                if target_filename not in description:
                    description = f"Target file '{target_filename}'. " + description

                if verbose:
                    print(f"📝 Sub-task {i+1}/{len(subtasks)}: [{task_type}]")
                    print(f"   Description: {description[:100]}...")

                if task_type == "FILE_MANIPULATION":
                    result, sub_tokens = self._handle_file_manipulation(
                        description=description,
                        target_filename=target_filename,
                        is_first_subtask=is_first_subtask,
                        mode=mode,
                        verbose=verbose,
                    )
                elif task_type == "APP_CONTROL":
                    result, sub_tokens = self._handle_app_control(
                        description=description,
                        target_filename=target_filename,
                        verbose=verbose,
                    )
                else:
                    result = {
                        "success": False,
                        "error": f"Unknown task type: {task_type}",
                        "output": "",
                    }
                    sub_tokens = 0

                total_tokens += sub_tokens

                results["subtasks"].append({
                    "type": task_type,
                    "description": description,
                    "result": result,
                })

                if result.get("success", False):
                    any_succeeded = True
                else:
                    all_succeeded = False
                    err = result.get("error", "Unknown error")
                    results["errors"].append(err)

            if all_succeeded:
                results["status"] = "success"
            elif any_succeeded:
                results["status"] = "partial"
            else:
                results["status"] = "failure"

        except Exception as e:
            results["status"] = "failure"
            results["errors"].append(str(e))

        results["total_tokens_used"] = total_tokens
        results["total_tokens"] = total_tokens
        self.execution_log.append(dict(results))

        if verbose:
            print(f"\n✅ Pipeline Finished. Status: {results['status'].upper()}")
            print(f"   Total Tokens: ~{total_tokens}")
            if results["errors"]:
                print("⚠️ Errors encountered:")
                for err in results["errors"]:
                    print(f"   - {err}")

        return results

    def _handle_file_manipulation(
        self,
        description: str,
        target_filename: str,
        is_first_subtask: bool,
        mode: str,
        verbose: bool,
    ) -> tuple[Dict[str, Any], int]:
        total_tokens = 0
        last_error = ""
        last_code = ""
        last_output = ""

        # Context guidance based on whether we are editing an existing document or creating new
        if mode == "EDIT_EXISTING":
            context_guidance = (
                f"EDITING EXISTING DOCUMENT: '{target_filename}' already exists and contains data. "
                f"Open it with `filename = '{target_filename}'; doc = Document(filename)`. "
                f"Preserve all existing paragraphs/content unless explicitly instructed to replace or modify them. "
                f"Save changes to `doc.save(filename)`."
            )
        elif is_first_subtask:
            context_guidance = (
                f"CREATING NEW DOCUMENT: Initialize a new document: `filename = '{target_filename}'; doc = Document()`. "
                f"Save to `doc.save(filename)`."
            )
        else:
            context_guidance = (
                f"LAYERING SUBTASK: Open the document created in subtask 1: `filename = '{target_filename}'; doc = Document(filename)`. "
                f"Do NOT re-add the main document title. Save to `doc.save(filename)`."
            )

        for attempt in range(1, self.max_retries + 1):
            try:
                if verbose:
                    retry_label = f" (Attempt {attempt}/{self.max_retries})" if attempt > 1 else ""
                    print(f"   📄 Generating python-docx code via Nemotron{retry_label}...")

                full_desc = f"{context_guidance}\n{description}"
                code, tokens = self.docx_gen.generate_with_usage(
                    full_desc,
                    self.nemotron,
                    model=self.config.nemotron_model,
                    error_context=last_error if attempt > 1 else None,
                )
                total_tokens += tokens
                last_code = code

                if verbose:
                    print(f"   Code generated ({len(code)} chars)")
                    print("   ⚙️ Executing python-docx code...")

                success, output, error = self.docx_gen.execute(code)
                last_output = output

                if success:
                    return {
                        "success": True,
                        "code": code,
                        "output": output,
                        "error": "",
                    }, total_tokens
                else:
                    last_error = error
                    if verbose:
                        print(f"   ⚠️ Execution attempt {attempt} failed: {error[:150]}...")

            except Exception as e:
                last_error = str(e)
                if verbose:
                    print(f"   ⚠️ Exception on attempt {attempt}: {last_error}")

        return {
            "success": False,
            "code": last_code,
            "output": last_output,
            "error": last_error,
        }, total_tokens

    def _handle_app_control(
        self,
        description: str,
        target_filename: str,
        verbose: bool,
    ) -> tuple[Dict[str, Any], int]:
        total_tokens = 0
        last_error = ""
        last_script = ""
        last_output = ""

        full_desc = f"Target document: '{target_filename}'.\n{description}"

        for attempt in range(1, self.max_retries + 1):
            try:
                if verbose:
                    retry_label = f" (Attempt {attempt}/{self.max_retries})" if attempt > 1 else ""
                    print(f"   🍎 Generating AppleScript via Osaurus-8B (or Nemotron fallback){retry_label}...")

                script, tokens = self.applescript_gen.generate_with_usage(
                    full_desc,
                    include_sdef=True,
                    model=self.config.osaurus_model,
                    error_context=last_error if attempt > 1 else None,
                )
                total_tokens += tokens
                last_script = script

                if verbose:
                    print(f"   Script generated ({len(script)} chars)")
                    print("   ✓ Verifying syntax with osacompile...")

                is_valid, compile_error = self.applescript_gen.compile_check(script)

                if not is_valid:
                    last_error = f"Compilation failed: {compile_error}"
                    if verbose:
                        print(f"   ❌ Syntax compilation failed: {compile_error[:150]}...")
                    continue

                if verbose:
                    print("   ⚙️ Executing AppleScript via osascript...")

                success, stdout, stderr = self.applescript_gen.execute(script)
                last_output = stdout

                if success:
                    return {
                        "success": True,
                        "script": script,
                        "output": stdout,
                        "error": stderr,
                    }, total_tokens
                else:
                    last_error = stderr
                    if verbose:
                        print(f"   ⚠️ osascript execution failed: {stderr[:150]}...")

            except Exception as e:
                last_error = str(e)

        return {
            "success": False,
            "script": last_script,
            "output": last_output,
            "error": last_error,
        }, total_tokens
